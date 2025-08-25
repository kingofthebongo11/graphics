"""Формирование отчета в формате DOCX."""

from __future__ import annotations

from pathlib import Path
import re
from docx import Document


def _format_section_title(folder_name: str) -> str:
    """Форматирует название топ-папки для заголовка раздела."""
    clean_name = re.sub(r"^\d+[\-_ ]*", "", folder_name)
    first, *rest = clean_name.split("-", 1)
    if rest:
        return f"{first} ({rest[0]})"
    return first


def analysis_prefix(path: Path) -> int:
    """Возвращает числовой префикс имени родительской папки PNG.

    Извлекает ведущие цифры из имени директории, содержащей ``path``.
    Если число не найдено, возвращает ``float('inf')``.

    Примеры:
        >>> analysis_prefix(Path('03-анализ/plot.png'))
        3
    """
    match = re.match(r"^(\d+)", path.parent.name)
    return int(match.group(1)) if match else float("inf")


def build_report(curves_root: Path | str) -> Path:
    """Собирает отчет по PNG файлам в каталоге Curves.

    Для каждой топ-папки создаёт заголовок первого уровня и добавляет
    все найденные в ней графики (PNG) из вложенных каталогов.
    Графики сортируются по числовому префиксу имени их родительской папки.

    Параметры:
        curves_root: Каталог с топ-папками.

    Возвращает:
        Путь к созданному файлу ``Report.docx``.
    """
    root = Path(curves_root)
    document = Document()

    def _prefix(path: Path) -> int:
        match = re.match(r"^(\d+)", path.name)
        return int(match.group(1)) if match else float("inf")

    for top_path in sorted(filter(Path.is_dir, root.iterdir()), key=_prefix):
        document.add_heading(_format_section_title(top_path.name), level=1)
        for image_path in sorted(top_path.rglob("*.png"), key=lambda p: analysis_prefix(p)):
            document.add_picture(str(image_path))

    output_path = root / "Report.docx"
    document.save(str(output_path))
    return output_path


if __name__ == "__main__":
    build_report(Path(__file__).resolve().parent / "Curves")
